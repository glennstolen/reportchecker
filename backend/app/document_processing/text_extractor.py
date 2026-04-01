import io
import re
from dataclasses import dataclass
from datetime import date
import fitz  # PyMuPDF
from docx import Document


@dataclass
class ReportMetadata:
    """Metadata extracted from report first page."""
    kandidater: list[int]  # List of candidate numbers
    oppgave: str | None
    dato: date | None


def extract_metadata_from_text(text: str) -> ReportMetadata:
    """Extract kandidat, oppgave, and dato from text (typically first page)."""
    kandidater = []
    oppgave = None
    dato = None

    # Extract kandidat number(s) - look for "Kandidat:" or "Kandidatnummer:" followed by numbers
    # Can be multiple candidates separated by comma, "og", or "&"
    kandidat_pattern = r'[Kk]andidat(?:nummer)?[:\s]+([0-9,\s&og]+)'
    kandidat_match = re.search(kandidat_pattern, text)
    if kandidat_match:
        kandidat_str = kandidat_match.group(1)
        # Extract all numbers from the match
        numbers = re.findall(r'\d+', kandidat_str)
        kandidater = [int(n) for n in numbers]

    # Extract oppgave - look for "Oppgave:" followed by text until newline
    oppgave_pattern = r'[Oo]ppgave[:\s]+([^\n\r]+)'
    oppgave_match = re.search(oppgave_pattern, text)
    if oppgave_match:
        oppgave = oppgave_match.group(1).strip()

    # Extract dato - look for "Dato:" followed by a date
    # Support formats: DD.MM.YYYY, DD/MM/YYYY, DD-MM-YYYY, YYYY-MM-DD, "28. mars 2026"
    dato_pattern = r'[Dd]ato[:\s]+(\d{1,2}\.?\s*[a-zæøåA-ZÆØÅ]+\.?\s*\d{4}|\d{1,2}[./-]\d{1,2}[./-]\d{2,4}|\d{4}[./-]\d{1,2}[./-]\d{1,2})'
    dato_match = re.search(dato_pattern, text)
    if dato_match:
        dato_str = dato_match.group(1)
        dato = parse_date(dato_str)

    return ReportMetadata(kandidater=kandidater, oppgave=oppgave, dato=dato)


NORWEGIAN_MONTHS = {
    'januar': 1, 'jan': 1,
    'februar': 2, 'feb': 2,
    'mars': 3, 'mar': 3,
    'april': 4, 'apr': 4,
    'mai': 5,
    'juni': 6, 'jun': 6,
    'juli': 7, 'jul': 7,
    'august': 8, 'aug': 8,
    'september': 9, 'sep': 9, 'sept': 9,
    'oktober': 10, 'okt': 10,
    'november': 11, 'nov': 11,
    'desember': 12, 'des': 12,
}


def parse_date(date_str: str) -> date | None:
    """Parse various date formats including Norwegian text dates."""
    # Clean up the string
    date_str = date_str.strip()

    # Try Norwegian text format: "28. mars 2026" or "28 mars 2026"
    text_pattern = r'(\d{1,2})\.?\s*([a-zæøå]+)\.?\s*(\d{4})'
    text_match = re.match(text_pattern, date_str, re.IGNORECASE)
    if text_match:
        day = int(text_match.group(1))
        month_str = text_match.group(2).lower()
        year = int(text_match.group(3))
        if month_str in NORWEGIAN_MONTHS:
            try:
                return date(year, NORWEGIAN_MONTHS[month_str], day)
            except ValueError:
                pass

    # Try numeric formats
    formats = [
        (r'(\d{1,2})[./-](\d{1,2})[./-](\d{4})', 'dmy'),  # DD.MM.YYYY
        (r'(\d{1,2})[./-](\d{1,2})[./-](\d{2})', 'dmy2'),  # DD.MM.YY
        (r'(\d{4})[./-](\d{1,2})[./-](\d{1,2})', 'ymd'),  # YYYY-MM-DD
    ]

    for pattern, fmt in formats:
        match = re.match(pattern, date_str)
        if match:
            try:
                if fmt == 'dmy':
                    return date(int(match.group(3)), int(match.group(2)), int(match.group(1)))
                elif fmt == 'dmy2':
                    year = int(match.group(3))
                    year = 2000 + year if year < 100 else year
                    return date(year, int(match.group(2)), int(match.group(1)))
                elif fmt == 'ymd':
                    return date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
            except ValueError:
                continue

    return None


def extract_first_page_text(content: bytes, file_extension: str) -> str:
    """Extract text from the first page only."""
    file_extension = file_extension.lower().lstrip(".")

    if file_extension == "pdf":
        with fitz.open(stream=content, filetype="pdf") as doc:
            if len(doc) > 0:
                return doc[0].get_text()
        return ""
    elif file_extension in ("docx", "doc"):
        # For Word docs, take first ~1000 characters as "first page"
        full_text = extract_text_from_docx(content)
        return full_text[:2000]
    return ""


def extract_text_from_file(content: bytes, file_extension: str) -> str:
    """Extract text content from a document file."""
    file_extension = file_extension.lower().lstrip(".")

    if file_extension == "pdf":
        return extract_text_from_pdf(content)
    elif file_extension in ("docx", "doc"):
        return extract_text_from_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from a PDF file."""
    text_parts = []

    with fitz.open(stream=content, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- Side {page_num} ---\n{page_text}")

    return "\n\n".join(text_parts)


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from a Word document."""
    doc = Document(io.BytesIO(content))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n\n".join(text_parts)
